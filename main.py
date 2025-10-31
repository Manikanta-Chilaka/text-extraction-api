from flask import Flask, request, jsonify
from flask_cors import CORS
import requests
import fitz  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX
import io
import os
from supabase import create_client, Client
import logging
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Supabase setup
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
    logger.warning("Supabase environment variables not set. Database updates will be disabled.")
    supabase = None
else:
    try:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
        logger.info("Supabase client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Supabase: {e}")
        supabase = None

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text += f"--- Page {page_num} ---\n{page_text}\n\n"
        doc.close()
        
        if not text.strip():
            raise ValueError("No text found in PDF. PDF might be image-based.")
        
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            raise ValueError("No text found in DOCX file.")
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file with multiple encoding attempts"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding).strip()
            if text:
                return text
        except (UnicodeDecodeError, AttributeError):
            continue
    
    raise Exception("Failed to decode text file with supported encodings")

def update_supabase_record(song_id: str, text: str):
    """Update Supabase record"""
    if not supabase:
        logger.warning("Supabase not configured, skipping database update")
        return False
    
    try:
        result = supabase.table('song_scripts').update({
            'content': text,
            'updated_at': datetime.utcnow().isoformat()
        }).eq('id', song_id).execute()
        
        if result.data:
            logger.info(f"Successfully updated song_id: {song_id}")
            return True
        else:
            logger.warning(f"No record found for song_id: {song_id}")
            return False
    except Exception as e:
        logger.error(f"Supabase update error: {e}")
        return False

@app.route('/extract-text', methods=['POST'])
def extract_text():
    """Extract text from uploaded file and optionally update Supabase"""
    try:
        data = request.get_json()
        if not data or 'file_url' not in data:
            return jsonify({'error': 'file_url is required'}), 400
        
        file_url = data['file_url']
        song_id = data.get('song_id')
        
        logger.info(f"Processing file: {file_url}")
        
        # Download file from URL
        response = requests.get(file_url, timeout=60)
        response.raise_for_status()
        file_content = response.content
        
        if not file_content:
            return jsonify({'error': 'Downloaded file is empty'}), 400
        
        # Determine file type and extract text
        file_url_lower = file_url.lower()
        
        if file_url_lower.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
            file_type = "pdf"
        elif file_url_lower.endswith('.docx'):
            text = extract_text_from_docx(file_content)
            file_type = "docx"
        elif file_url_lower.endswith('.txt'):
            text = extract_text_from_txt(file_content)
            file_type = "txt"
        else:
            return jsonify({
                'error': 'Unsupported file type. Supported formats: PDF, DOCX, TXT'
            }), 400
        
        # Update Supabase if song_id provided
        updated_supabase = False
        if song_id and supabase:
            updated_supabase = update_supabase_record(song_id, text)
        
        logger.info(f"Successfully extracted {len(text)} characters from {file_type}")
        
        return jsonify({
            'text': text,
            'status': 'success',
            'file_type': file_type,
            'updated_supabase': updated_supabase,
            'char_count': len(text),
            'timestamp': datetime.utcnow().isoformat()
        })
        
    except requests.RequestException as e:
        logger.error(f"File download error: {e}")
        return jsonify({'error': f'Failed to download file: {str(e)}'}), 400
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return jsonify({'error': f'Extraction failed: {str(e)}'}), 500

@app.route('/')
def root():
    return jsonify({
        "service": "Church App Text Extraction API",
        "version": "1.0.0",
        "status": "operational",
        "supported_formats": ["PDF", "DOCX", "TXT"],
        "endpoints": {
            "POST /extract-text": "Extract text from file URL",
            "GET /health": "Health check endpoint"
        }
    })

@app.route('/health')
def health_check():
    """Health check endpoint for monitoring"""
    supabase_status = "connected" if supabase else "disabled"
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "supabase": supabase_status
    })

if __name__ == "__main__":
    port = int(os.getenv("PORT", 8000))
    app.run(host="0.0.0.0", port=port, debug=False)